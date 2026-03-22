import tempfile
from pathlib import Path

from docx import Document
from fpdf import FPDF
from openpyxl import Workbook


def format_text_response(answer: str) -> str:
    return answer


def create_text_file(answer: str) -> Path:
    with tempfile.NamedTemporaryFile(
        mode="w",
        encoding="utf-8",
        suffix=".txt",
        prefix="moodle_ai_",
        delete=False,
    ) as temp_file:
        temp_file.write(answer)
        return Path(temp_file.name)


def create_pdf(answer: str) -> Path:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Helvetica", size=12)

    for line in answer.splitlines() or [answer]:
        pdf.multi_cell(0, 8, line)

    with tempfile.NamedTemporaryFile(suffix=".pdf", prefix="moodle_ai_", delete=False) as temp_file:
        path = temp_file.name
    pdf.output(path)
    return Path(path)


def create_excel(answer: str) -> Path:
    wb = Workbook()
    ws = wb.active
    ws.title = "Response"
    ws.cell(row=1, column=1, value="Moodle AI Assistant Response")

    row = 3
    for line in answer.splitlines() or [answer]:
        ws.cell(row=row, column=1, value=line)
        row += 1

    with tempfile.NamedTemporaryFile(suffix=".xlsx", prefix="moodle_ai_", delete=False) as temp_file:
        path = temp_file.name
    wb.save(path)
    return Path(path)


def create_word(answer: str) -> Path:
    document = Document()
    document.add_heading("Moodle AI Assistant Response", level=1)

    for line in answer.splitlines() or [answer]:
        document.add_paragraph(line)

    with tempfile.NamedTemporaryFile(suffix=".docx", prefix="moodle_ai_", delete=False) as temp_file:
        path = temp_file.name
    document.save(path)
    return Path(path)
